from flask import Flask, request, render_template, jsonify
from werkzeug.utils import secure_filename
import os

from HybridRetriever import HybridRetriever
from ChatEngine import ChatEngine
from configs import *

import csv
import fitz
from docx import Document as DocxDocument

import logging
import warnings
warnings.simplefilter("ignore", UserWarning)
warnings.simplefilter("ignore", FutureWarning)

from llama_index.retrievers.bm25 import BM25Retriever 
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from vllm import LLM

def process_file(file) -> Document:
    """
    Processes the given file and extracts its text content to create a llama-index Document.

    This function supports various file formats including:
    - Text files (.txt)
    - CSV files (.csv)
    - PDF files (.pdf)
    - DOCX files (.docx)

    Based on the file extension, the function reads the content of the file accordingly:
    - For text files, it reads the entire file content as a string.
    - For CSV files, it reads each row, joining them into a single string with new lines.
    - For PDF files, it extracts text from each page and concatenates them.
    - For DOCX files, it reads each paragraph and combines them into a single string.

    Parameters:
        file (str): The path to the file to be processed.

    Returns:
        Document: A llama-index Document containing the extracted text content.
    """
    logging.info(f"Processing file: {file}")
    file_extension = file.split(".")[-1].lower()

    if file_extension == 'txt':
        with open(file, 'r', encoding='utf-8') as f:
            text = f.read()
            logging.info("Processed text file")

    elif file_extension == 'csv':
        with open(file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = '\n'.join(','.join(row) for row in reader)
            logging.info("Processed CSV file")

    elif file_extension == 'pdf':
        pdf_document = fitz.open(file, filetype=file_extension)
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        pdf_document.close()
        logging.info("Processed PDF file")
        
    elif file_extension == 'docx':
        docx_document = DocxDocument(file)
        text = ""
        for paragraph in docx_document.paragraphs:
            text += paragraph.text + "\n"
        logging.info("Processed DOCX file")

    return [Document(text=text)]

def process_and_respond(file, question) -> str:
    """
    Processes the given file, indexes its content, and generates a response to a question using a chat engine.

    This function performs the following steps:
    1. Processes the file to extract its text content using the `process_file` function.
    2. Splits the extracted text into manageable chunks using a sentence splitter.
    3. Creates a vector index from the text chunks and initializes a BM25 and vector retriever.
    4. Combines the retrievers into a hybrid retriever for improved search capabilities.
    5. Initializes a chat engine using the hybrid retriever.
    6. Generates a response to the provided question using the chat engine and the global language model (`llm`).

    Parameters:
        file (str): The path to the file to be processed.
        question (str): The question to be answered based on the file's content.

    Returns:
        response: The generated response to the question.
    """
    global llm
    logging.info(f"Starting to process file: {file}")
    documents = process_file(file)
    
    logging.info("Splitting text into chunks")
    text_splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    logging.info("Creating vector index from documents")
    vector_index = VectorStoreIndex.from_documents(
        documents, transformations=[text_splitter], embed_model=Settings.embed_model, show_progress=True
    )

    logging.info("Initializing retrievers")
    bm25_retriever = BM25Retriever(nodes=documents, similarity_top_k=TOP_K, tokenizer=text_splitter.split_text)
    vector_retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=TOP_K)
    hybrid_retriever = HybridRetriever(bm25_retriever=bm25_retriever, vector_retriever=vector_retriever)

    logging.info("Initializing chat engine")
    chat_engine = ChatEngine(hybrid_retriever)

    logging.info("Generating response")
    response = chat_engine.ask_question(question, llm)
    
    logging.info("Response generated")
    logging.info(response)
    return response


if __name__ == '__main__':
    app = Flask(__name__)
    app.config['UPLOAD_FOLDER'] = 'uploads'

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

    logging.basicConfig(level=logging.INFO)
    
    llm = LLM(MODEL_NAME, tensor_parallel_size=4, dtype="half")
    embedding = HuggingFaceEmbedding(
        model_name=EMBEDDING_NAME,
        device="cuda:2",
        trust_remote_code=True,
        )
    
    Settings.llm = llm
    Settings.embed_model = embedding
    logging.info("Initialized LLM and embedding models")
    
    @app.route('/')
    def home():
        logging.info("Rendering home page")
        return render_template('index.html')

    @app.route('/upload', methods=['POST'])
    def upload_file():
        logging.info("Received file upload request")
        
        if 'file' not in request.files:
            logging.error("No file part in the request")
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']

        if file.filename == '':
            logging.error("No selected file")
            return jsonify({'error': 'No selected file'}), 400

        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            logging.info(f"File saved to {file_path}")
            
            question = request.form.get('question', '')
            logging.info(f"Processing file with question: {question}")
            
            response = process_and_respond(file_path, question)
            
            os.remove(file_path)
            logging.info(f"File {file_path} removed")
            
            return {"response": response} 
        
    logging.info("Starting Flask app")
    app.run(host='0.0.0.0', port=5000)
