

from HybridRetriever import HybridRetriever
from ChatEngine import ChatEngine
from configs import *
import csv
import warnings
warnings.simplefilter("ignore", UserWarning)
warnings.simplefilter("ignore", FutureWarning)

import logging
logging.basicConfig(level=logging.INFO)

from llama_index.retrievers.bm25 import BM25Retriever 
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core import VectorStoreIndex, Document
from llama_index.core import Settings
from llama_index.core.node_parser import SentenceSplitter
import fitz
from docx import Document as DocxDocument

from transformers import AutoTokenizer, AutoModelForCausalLM
from transformers import BitsAndBytesConfig
import torch
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

#quantization_config = BitsAndBytesConfig(
#        load_in_4bit=True,
#        bnb_4bit_compute_dtype=torch.float16,
#        bnb_4bit_quant_type="nf4",
#        bnb_4bit_use_double_quant=True,
#        )
from transformers import AutoTokenizer, AutoModelForCausalLM
import deepspeed

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForCausalLM.from_pretrained(MODEL_NAME, torch_dtype="auto", trust_remote_code=True)

# init deepspeed inference engine
llm = deepspeed.init_inference(
    model=model,      # Transformers models
    mp_size=4,        # Number of GPU
    dtype=torch.half, # dtype of the weights (fp16)
    # injection_policy={"BertLayer" : HFBertLayerPolicy}, # replace BertLayer with DS HFBertLayerPolicy
    #replace_method="auto", # Lets DS autmatically identify the layer to replace
    replace_with_kernel_inject=True, # replace the model with the kernel injector
    max_tokens=2048,
)

#llm = HuggingFaceLLM(
#        model_name=MODEL_NAME,
#        tokenizer_name=MODEL_NAME,
#        context_window=CONTEXT_WINDOW,
#        #model_kwargs={"quantization_config": quantization_config},
#        generate_kwargs={"temperature": TEMPERATURE},
#        device_map=DEVICE,
#        )   
embedding = HuggingFaceEmbedding(
        model_name=EMBEDDING_NAME,
        device="cuda:2",
        trust_remote_code=True,
        )

logging.info("Initializing LLM and embedding models")
Settings.llm = llm
Settings.embed_model = embedding

def process_file(file):
    logging.info(f"Processing file: {file}")
    file_extension = file.split(".")[-1].lower()

    if file_extension == 'txt':
        with open(file, 'r', encoding='utf-8') as f:
            text = f.read()
            logging.info("Processed text file")

    elif file_extension == 'csv':
        with open(file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = '\n'.join(','.join(row) for row in reader)
            logging.info("Processed CSV file")

    elif file_extension == 'pdf':
        pdf_document = fitz.open(file, filetype=file_extension)
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        pdf_document.close()
        logging.info("Processed PDF file")
        
    elif file_extension == 'docx':
        docx_document = DocxDocument(file)
        text = ""
        for paragraph in docx_document.paragraphs:
            text += paragraph.text + "\n"
            logging.info("Processed DOCX file")

    return [Document(text=text)]

def process_and_respond(file, question):
    global llm
    logging.info(f"Starting to process file: {file}")
    documents = process_file(file)
    
    logging.info("Splitting text into chunks")
    text_splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    logging.info("Creating vector index from documents")
    vector_index = VectorStoreIndex.from_documents(
        documents, transformations=[text_splitter], embed_model=Settings.embed_model, show_progress=True
    )
    logging.info("Initializing retrievers")
    bm25_retriever = BM25Retriever(nodes=documents, similarity_top_k=TOP_K, tokenizer=text_splitter.split_text)
    vector_retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=TOP_K)
    hybrid_retriever = HybridRetriever(bm25_retriever=bm25_retriever, vector_retriever=vector_retriever)
    logging.info("Initializing chat engine")
    chat_engine = ChatEngine(hybrid_retriever)
    logging.info("Generating response")
    response = chat_engine.ask_question(question, llm)
    logging.info("Response generated")
    return response
